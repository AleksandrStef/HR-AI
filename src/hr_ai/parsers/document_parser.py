"""
Document parser module for extracting text from various document formats.
Supports .docx, .doc, and .pdf files commonly used for IDPs.
"""

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import logging

from docx import Document
import PyPDF2
from docx.shared import Inches

logger = logging.getLogger(__name__)

class DocumentParseError(Exception):
    """Custom exception for document parsing errors."""
    pass

class DocumentParser:
    """Parser for extracting text and metadata from IDP documents."""
    
    def __init__(self, docs_directory: str = "docs"):
        self.docs_directory = Path(docs_directory)
        self.supported_extensions = {'.docx', '.doc', '.pdf'}
        
    def parse_document(self, file_path: str) -> Dict[str, any]:
        """
        Parse a document and extract text content with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing extracted text, metadata, and structure information
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DocumentParseError(f"File not found: {file_path}")
            
        if file_path.suffix.lower() not in self.supported_extensions:
            raise DocumentParseError(f"Unsupported file format: {file_path.suffix}")
            
        try:
            if file_path.suffix.lower() == '.docx':
                return self._parse_docx(file_path)
            elif file_path.suffix.lower() == '.pdf':
                return self._parse_pdf(file_path)
            else:
                raise DocumentParseError(f"Parser not implemented for: {file_path.suffix}")
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise DocumentParseError(f"Failed to parse {file_path}: {str(e)}")
    
    def _parse_docx(self, file_path: Path) -> Dict[str, any]:
        """Parse a DOCX file and extract structured content."""
        doc = Document(str(file_path))
        
        # Extract basic metadata
        employee_name = self._extract_employee_name(file_path.name)
        
        # Extract all text content
        full_text = []
        sections = {}
        current_section = "intro"
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            full_text.append(text)
            
            # Detect section headers
            section_header = self._detect_section_header(text)
            if section_header:
                current_section = section_header
                sections[current_section] = []
            else:
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(text)
        
        # Extract tables if any
        tables_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Skip empty rows
                    table_data.append(row_data)
            if table_data:
                tables_content.append(table_data)
        
        # Extract dates and meeting information
        dates_found = self._extract_dates(full_text)
        meeting_sections = self._identify_meeting_sections(sections)
        
        return {
            'file_path': str(file_path),
            'employee_name': employee_name,
            'full_text': '\n'.join(full_text),
            'sections': sections,
            'tables': tables_content,
            'dates_found': dates_found,
            'meeting_sections': meeting_sections,
            'word_count': len(' '.join(full_text).split()),
            'parsed_at': datetime.now().isoformat(),
            'file_modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
        }
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, any]:
        """Parse a PDF file and extract text content."""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            full_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    full_text.append(text)
        
        text_content = '\n'.join(full_text)
        employee_name = self._extract_employee_name(file_path.name)
        dates_found = self._extract_dates([text_content])
        
        return {
            'file_path': str(file_path),
            'employee_name': employee_name,
            'full_text': text_content,
            'sections': {'full_document': [text_content]},
            'tables': [],
            'dates_found': dates_found,
            'meeting_sections': [],
            'word_count': len(text_content.split()),
            'parsed_at': datetime.now().isoformat(),
            'file_modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
        }
    
    def _extract_employee_name(self, filename: str) -> str:
        """Extract employee name from filename."""
        # Remove extension and common phrases
        name = Path(filename).stem
        name = re.sub(r'(Employee development plan|План развития сотрудника)', '', name, flags=re.IGNORECASE)
        name = re.sub(r'^-\s*|\s*-$', '', name)  # Remove leading/trailing dashes
        return name.strip()
    
    def _detect_section_header(self, text: str) -> Optional[str]:
        """Detect if text is a section header and return normalized section name."""
        text_lower = text.lower()
        
        # Common section patterns
        section_patterns = {
            'plans_before_review': [
                'plans before', 'планы до ревью', 'планируется',
                'probation period', 'испытательный срок'
            ],
            'performance_review': [
                'performance review', 'годовое ревью', 'annual review'
            ],
            'quarterly_checkpoint': [
                'quarterly', 'checkpoint', 'чек-поинт', 'квартальный'
            ],
            'goals': [
                'goals for', 'цели на', 'targets', 'objectives'
            ],
            'feedback': [
                'feedback', 'обратная связь', 'что нравится', 'what do you like'
            ],
            'satisfaction': [
                'satisfaction', 'удовлетворен', 'отношение к компании'
            ],
            'training': [
                'training', 'обучение', 'certification', 'сертификация'
            ],
            'location': [
                'location', 'локация', 'relocation', 'релокация'
            ]
        }
        
        for section_key, patterns in section_patterns.items():
            if any(pattern in text_lower for pattern in patterns):
                return section_key
        
        return None
    
    def _extract_dates(self, text_lines: List[str]) -> List[Dict[str, str]]:
        """Extract dates from text content."""
        dates = []
        date_patterns = [
            r'\d{1,2}[./]\d{1,2}[./]\d{2,4}',  # DD/MM/YYYY or DD.MM.YYYY
            r'\d{2,4}[-/]\d{1,2}[-/]\d{1,2}',  # YYYY-MM-DD
            r'\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',
            r'\d{1,2}\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4}'
        ]
        
        for line in text_lines:
            for pattern in date_patterns:
                matches = re.finditer(pattern, line, re.IGNORECASE)
                for match in matches:
                    dates.append({
                        'date_string': match.group(),
                        'context': line.strip(),
                        'position': match.span()
                    })
        
        return dates
    
    def _identify_meeting_sections(self, sections: Dict[str, List[str]]) -> List[str]:
        """Identify which sections likely contain meeting information."""
        meeting_sections = []
        
        meeting_indicators = [
            'checkpoint', 'review', 'meeting', 'встреча', 'обсуждение',
            'созвон', 'беседа', 'разговор'
        ]
        
        for section_name, content in sections.items():
            section_text = ' '.join(content).lower()
            if any(indicator in section_text for indicator in meeting_indicators):
                meeting_sections.append(section_name)
        
        return meeting_sections
    
    def scan_directory(self) -> List[Dict[str, any]]:
        """
        Scan the documents directory and return information about all files.
        
        Returns:
            List of dictionaries containing file information
        """
        files_info = []
        
        if not self.docs_directory.exists():
            logger.warning(f"Documents directory does not exist: {self.docs_directory}")
            return files_info
        
        for file_path in self.docs_directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    stat = file_path.stat()
                    files_info.append({
                        'file_path': str(file_path),
                        'employee_name': self._extract_employee_name(file_path.name),
                        'file_size': stat.st_size,
                        'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        'extension': file_path.suffix.lower()
                    })
                except Exception as e:
                    logger.warning(f"Could not get info for {file_path}: {str(e)}")
        
        return files_info
    
    def get_recently_modified_files(self, days: int = 7) -> List[str]:
        """
        Get list of files modified within the specified number of days.
        
        Args:
            days: Number of days to look back
            
        Returns:
            List of file paths
        """
        cutoff_time = datetime.now().timestamp() - (days * 24 * 60 * 60)
        files_info = self.scan_directory()
        
        recent_files = []
        for file_info in files_info:
            modified_timestamp = datetime.fromisoformat(file_info['modified_time']).timestamp()
            if modified_timestamp >= cutoff_time:
                recent_files.append(file_info['file_path'])
        
        return recent_files